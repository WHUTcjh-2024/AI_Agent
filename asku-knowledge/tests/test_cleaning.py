import hashlib
import io
import sqlite3
import tempfile
import unittest
import zipfile
from contextlib import closing
from pathlib import Path

import yaml
from asku.admission import evaluate, official_url, text_hash
from asku.batch_validation import finalize_and_verify
from asku.copy_dump import read_copy
from asku.document_parser import markdown_table, parse_bytes
from asku.local_cleaning import rebuild_links, run_batch, safe_raw_path
from asku.normalizer import normalize_html

ROOT = Path(__file__).resolve().parents[2]


class AdmissionTests(unittest.TestCase):
    def setUp(self):
        self.school = yaml.safe_load(
            (ROOT / "config/schools/whut.yaml").read_text(encoding="utf-8")
        )
        self.taxonomy = yaml.safe_load(
            (ROOT / "asku-knowledge/config/taxonomy.yaml").read_text(encoding="utf-8")
        )
        self.document = dict(
            school_id="whut",
            review_status="ACCEPTED",
            parse_status="PARSED",
            content_chars=200,
            normalized_path="normalized/d.md",
            content_hash="abc",
            pii_scan_status="CLEAR",
            pii_content_hash="abc",
            pii_detected=0,
            source_url="https://jwc.whut.edu.cn/notice/1.htm",
            canonical_url="https://jwc.whut.edu.cn/notice/1.htm",
            secondary_topic="recommendation",
            primary_module="UNDERGRAD_ACADEMIC",
            education_level="UNDERGRADUATE",
            audience="UNDERGRADUATE",
            document_type="POLICY_REGULATION",
            is_attachment=0,
        )

    def gate(self, **changes):
        return evaluate(
            {**self.document, **changes}, self.taxonomy, self.school, source_active=True
        )

    def test_valid_reviewed_text_is_candidate(self):
        self.assertTrue(self.gate().eligible)

    def test_flags_cannot_bypass_missing_parsing_or_review(self):
        for changes in [
            dict(parse_status="FAILED"),
            dict(review_status="REVIEW"),
            dict(content_chars=0),
            dict(normalized_path=""),
            dict(pii_scan_status="NOT_SCANNED"),
            dict(pii_detected=1),
            dict(pii_content_hash="stale"),
            dict(canonical_document_id="original"),
        ]:
            with self.subTest(changes=changes):
                self.assertFalse(self.gate(rag_eligible=1, **changes).eligible)

    def test_other_and_discovery_are_never_admitted(self):
        self.assertFalse(
            self.gate(secondary_topic="other", primary_module="OTHER").eligible
        )
        self.assertFalse(self.gate(document_type="INDEX_PAGE").eligible)

    def test_attachment_needs_evidenced_parent(self):
        self.assertFalse(
            self.gate(is_attachment=1, parent_page_url="whd_fake").eligible
        )
        self.assertTrue(
            self.gate(
                is_attachment=1,
                parent_page_url=self.document["source_url"],
                relation_status="RESOLVED",
                knowledge_bundle_id="b",
            ).eligible
        )

    def test_official_url_checks_host_and_forbidden_paths(self):
        for url in [
            "https://whut.edu.cn.evil.test/x",
            "https://evil.test/?x=whut.edu.cn",
            "https://u:p@jwc.whut.edu.cn/x",
            "https://mail.whut.edu.cn/x",
            "https://jwc.whut.edu.cn/login",
        ]:
            self.assertFalse(official_url(url, self.school), url)


class ParsingTests(unittest.TestCase):
    def test_binary_is_not_text_even_if_named_txt(self):
        result = parse_bytes(b"PK\x00\x00garbage" * 100, "policy.txt")
        self.assertEqual(result.status, "FAILED")
        self.assertEqual(result.text, "")

    def test_empty_text_is_not_success(self):
        self.assertEqual(parse_bytes(b"\n\n\t ", "empty.txt").status, "EMPTY")

    def test_blank_pdf_is_not_success(self):
        import pymupdf

        with pymupdf.open() as pdf:
            pdf.new_page()
            result = parse_bytes(pdf.tobytes(), "blank.pdf")
        self.assertEqual(result.status, "EMPTY")

    def test_docx_keeps_paragraph_table_order(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("申请条件：" + ("学生符合申请条件。" * 15))
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "材料"
        table.cell(0, 1).text = "份数"
        table.cell(1, 0).text = "申请表"
        table.cell(1, 1).text = "2"
        doc.add_paragraph("提交期限")
        buffer = io.BytesIO()
        doc.save(buffer)
        result = parse_bytes(buffer.getvalue(), "renamed.bin")
        self.assertEqual(result.status, "PARSED")
        self.assertLess(result.text.index("申请条件"), result.text.index("| 材料"))
        self.assertLess(result.text.index("| 材料"), result.text.index("提交期限"))

    def test_partial_archive_stays_review_and_paths_are_not_extracted(self):
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as z:
            z.writestr("../policy.txt", "申请程序说明。" * 30)
            z.writestr("broken.pdf", b"%PDF-broken")
        result = parse_bytes(buffer.getvalue(), "sample.zip")
        self.assertEqual(result.status, "REVIEW")
        self.assertEqual(len(result.members), 2)
        self.assertIn("申请程序说明", result.text)

    def test_table_escapes_cells(self):
        self.assertIn("a\\|b", markdown_table([["a|b", "c"], ["x\ny", "2"]]))

    def test_normalizer_counts_body_and_keeps_publish_metadata(self):
        body = "办理条件及申请程序。" * 25
        page = normalize_html(
            '<html><head><meta name="publishdate" content="2026-09-01"></head><body><div class="nav">'
            + ("菜单" * 200)
            + '</div><div class="TRS_Editor">'
            + body
            + "</div></body></html>"
        )
        self.assertEqual(page.text, body)
        self.assertEqual(page.content_chars, len(body))
        self.assertEqual(page.publish_date, "2026-09-01")

    def test_event_date_is_not_publication_date(self):
        page = normalize_html(
            "<html><body><main>会议时间：2026年10月22日。"
            + ("正文说明。" * 50)
            + "</main></body></html>"
        )
        self.assertIsNone(page.publish_date)


class LineageTests(unittest.TestCase):
    def test_copy_decoder_and_column_checks(self):
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "dump.sql"
            path.write_text(
                "COPY cck.raw_documents (id, raw_path) FROM stdin;\n1\tD:\\\\data\\\\a.html\n\\.\n",
                encoding="utf-8",
            )
            self.assertEqual(
                read_copy(path)["cck.raw_documents"][0]["raw_path"], r"D:\data\a.html"
            )
            path.write_text(
                "COPY cck.raw_documents (id, raw_path) FROM stdin;\n1\n\\.\n",
                encoding="utf-8",
            )
            with self.assertRaises(ValueError):
                read_copy(path)

    def test_raw_path_cannot_escape_configured_root(self):
        with tempfile.TemporaryDirectory() as folder:
            with self.assertRaises(ValueError):
                safe_raw_path("../outside.pdf", Path(folder))

    def test_relations_require_exact_anchor_not_filename(self):
        school = {"school_id": "whut", "allowed_domains": ["whut.edu.cn"]}
        with tempfile.TemporaryDirectory() as folder:
            raw = Path(folder)
            page = raw / "p.html"
            page.write_text(
                '<html><a href="../files/a.pdf">附件</a><a href="https://evil.test/files/b.pdf">伪附件</a></html>',
                encoding="utf-8",
            )
            docs = [
                dict(
                    id="p",
                    source_url="https://jwc.whut.edu.cn/n/p.html",
                    is_attachment=0,
                    raw_path=str(page),
                    title="通知",
                    source_id="s",
                ),
                dict(
                    id="a",
                    source_url="https://jwc.whut.edu.cn/files/a.pdf",
                    is_attachment=1,
                ),
                dict(
                    id="b",
                    source_url="https://jwc.whut.edu.cn/files/b.pdf",
                    is_attachment=1,
                ),
            ]
            relations, stats = rebuild_links(docs, {}, school, raw)
            self.assertEqual(relations, [("a", "p")])

    def test_text_fingerprint_normalizes_whitespace(self):
        self.assertEqual(text_hash("条件\n\n材料"), text_hash("条件 材料"))


class BatchTests(unittest.TestCase):
    def test_batch_preserves_source_and_validator_rejects_tampered_text(self):
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            raw = root / "raw"
            raw.mkdir()
            page = raw / "page.html"
            attachment = raw / "policy.txt"
            page.write_text(
                "<html><h1>课程申请流程</h1><main>"
                + ("申请需提交审批材料。" * 25)
                + '<a href="policy.txt">附件</a></main></html>',
                encoding="utf-8",
            )
            attachment.write_text(
                "学生应在指定时间提交申请材料。" * 30, encoding="utf-8"
            )
            base = dict(
                id="p",
                school_id="whut",
                source_id="s",
                title="课程申请流程",
                source_url="https://jwc.whut.edu.cn/page.html",
                canonical_url="",
                raw_path=str(page),
                normalized_path="",
                local_file_path="",
                publish_date=None,
                is_attachment=0,
                review_status="ACCEPTED",
                rag_eligible=1,
                pii_detected=0,
                pii_categories="[]",
                secondary_topic="course_selection",
                primary_module="UNDERGRAD_ACADEMIC",
                education_level="UNDERGRADUATE",
                audience="UNDERGRADUATE",
                document_type="PROCEDURE_GUIDE",
                content_chars=0,
                text_length=0,
                content_hash="",
                canonical_document_id=None,
                knowledge_bundle_id="",
                parent_page_url="",
                updated_at="",
                rejection_reason="",
            )
            source = root / "input.sqlite"
            with closing(sqlite3.connect(source)) as conn, conn:
                conn.execute(
                    "CREATE TABLE documents ("
                    + ",".join(
                        '"'
                        + k
                        + '" '
                        + ("INTEGER" if isinstance(v, int) else "TEXT")
                        + (" PRIMARY KEY" if k == "id" else "")
                        for k, v in base.items()
                    )
                    + ")"
                )
                for row in [
                    base,
                    {
                        **base,
                        "id": "a",
                        "is_attachment": 1,
                        "raw_path": str(attachment),
                        "source_url": "https://jwc.whut.edu.cn/policy.txt",
                    },
                ]:
                    conn.execute(
                        "INSERT INTO documents VALUES ("
                        + ",".join("?" for _ in row)
                        + ")",
                        list(row.values()),
                    )
                conn.execute(
                    "CREATE TABLE sources(id TEXT PRIMARY KEY,school_id TEXT,active INTEGER)"
                )
                conn.execute("INSERT INTO sources VALUES ('s','whut',1)")
                conn.execute(
                    "CREATE TABLE knowledge_bundles(id TEXT PRIMARY KEY,school_id TEXT,title TEXT,primary_document_id TEXT,document_count INTEGER,attachment_count INTEGER,created_at TEXT,updated_at TEXT)"
                )
                conn.execute(
                    "CREATE TABLE attachments(id TEXT PRIMARY KEY,document_id TEXT,parent_document_id TEXT,parent_page_url TEXT,knowledge_bundle_id TEXT,local_file_path TEXT,file_path TEXT,rag_eligible INTEGER,review_status TEXT,pii_detected INTEGER,updated_at TEXT)"
                )
                conn.execute(
                    "INSERT INTO attachments(id,document_id) VALUES ('attachment','a')"
                )
                conn.execute(
                    "CREATE TABLE weknora_mappings(school_id TEXT,import_status TEXT)"
                )
            before = hashlib.sha256(source.read_bytes()).hexdigest()
            dump = root / "crawler.sql"
            dump.write_text("", encoding="utf-8")
            output = root / "batch"
            legacy = root / "legacy"
            legacy.mkdir()
            school = ROOT / "config/schools/whut.yaml"
            taxonomy = ROOT / "asku-knowledge/config/taxonomy.yaml"
            summary = run_batch(
                source_db=source,
                crawler_dump=dump,
                school_path=school,
                taxonomy_path=taxonomy,
                raw_root=raw,
                output=output,
                legacy_dir=legacy,
                tesseract=None,
                workers=1,
            )
            self.assertEqual(summary["linked_attachments"], 1)
            self.assertEqual(summary["ready_documents"], 2)
            self.assertEqual(before, hashlib.sha256(source.read_bytes()).hexdigest())
            result = finalize_and_verify(output, school, taxonomy)
            self.assertEqual(result["status"], "PASSED")
            self.assertEqual(result["ready_attachments"], 1)
            (output / "normalized/p.md").write_text("篡改正文", encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "batch_validation_failed"):
                finalize_and_verify(output, school, taxonomy)


if __name__ == "__main__":
    unittest.main()
